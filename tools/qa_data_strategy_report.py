from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.oxml.ns import qn


path = Path(__file__).resolve().parents[1] / "deliverables" / "OY_Dashboard_Turkiye_Verisi_ve_AI_Altyapi_Strateji_Raporu.docx"

with ZipFile(path) as archive:
    assert archive.testzip() is None, "DOCX ZIP bütünlüğü bozuk"
    xml = archive.read("word/document.xml").decode("utf-8")
    rels = archive.read("word/_rels/document.xml.rels").decode("utf-8")
    footer_xml = "\n".join(
        archive.read(name).decode("utf-8")
        for name in archive.namelist()
        if name.startswith("word/footer") and name.endswith(".xml")
    )

doc = Document(path)
all_text = "\n".join(p.text for p in doc.paragraphs)
all_text += "\n" + "\n".join(cell.text for table in doc.tables for row in table.rows for cell in row.cells)

for forbidden in ("cite", "turn0search", "turn1search", "Lorem ipsum", "TODO", "AI tarafından oluştur"):
    assert forbidden not in all_text, f"Yasak/yer tutucu metin bulundu: {forbidden}"

h1 = sum(1 for p in doc.paragraphs if p.style and p.style.name == "Heading 1")
h2 = sum(1 for p in doc.paragraphs if p.style and p.style.name == "Heading 2")
assert h1 == 16, h1
assert h2 == 23, h2
assert len(doc.sections) == 2
assert len(doc.tables) == 17
assert xml.count("w:tblHeader") == 12, xml.count("w:tblHeader")
assert "Önerilen Türkiye İçi Veri ve AI Çalışma Mimarisi" not in all_text  # diagram is embedded, not duplicated as hidden text
assert rels.count("TargetMode=\"External\"") == 17, rels.count("TargetMode=\"External\"")
assert 'descr=""' not in xml
assert "hedef mimari diyagramı" in xml
assert " PAGE " in xml or " PAGE " in footer_xml

print(f"OK: {path}")
print(f"paragraphs={len(doc.paragraphs)} tables={len(doc.tables)} sections={len(doc.sections)} h1={h1} h2={h2}")
external_links = rels.count('TargetMode="External"')
print(f"external_links={external_links} size_bytes={path.stat().st_size}")
