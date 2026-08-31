from __future__ import annotations

from datetime import date
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import nsdecls, qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "deliverables"
TMP_DIR = ROOT / "tmp" / "data_strategy_report"
OUT_FILE = OUT_DIR / "OY_Dashboard_Turkiye_Verisi_ve_AI_Altyapi_Strateji_Raporu.docx"
DIAGRAM_FILE = TMP_DIR / "hedef_mimari.png"


# Resolved preset: standard_business_brief
PAGE_WIDTH_IN = 8.5
PAGE_HEIGHT_IN = 11.0
MARGIN_IN = 1.0
HEADER_FOOTER_IN = 0.492
CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGINS_DXA = {"top": 80, "bottom": 80, "start": 120, "end": 120}

NAVY = "0B2545"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "1F2937"
MUTED = "5B6573"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"
PALE_GREEN = "EAF4EE"
GREEN = "245B3A"
PALE_GOLD = "FFF6DD"
GOLD = "7A5A00"
PALE_RED = "FBECEC"
RED = "9B1C1C"
WHITE = "FFFFFF"
BORDER = "C7CFD8"


def rgb(hex_value: str) -> RGBColor:
    return RGBColor.from_string(hex_value)


def set_run_font(run, name="Calibri", size=None, color=INK, bold=None, italic=None):
    run.font.name = name
    if run._element.get_or_add_rPr().rFonts is None:
        rfonts = OxmlElement("w:rFonts")
        run._element.rPr.insert(0, rfonts)
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_borders(cell, color=BORDER, size="6"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = qn(f"w:{edge}")
        border = borders.find(tag)
        if border is None:
            border = OxmlElement(f"w:{edge}")
            borders.append(border)
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), size)
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), color)


def set_table_geometry(table, widths_dxa: list[int], indent_dxa=TABLE_INDENT_DXA):
    assert sum(widths_dxa) == CONTENT_WIDTH_DXA, (widths_dxa, sum(widths_dxa))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(CONTENT_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(widths_dxa[idx] / 1440)
            set_cell_margins(cell, **CELL_MARGINS_DXA)
            set_cell_borders(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_repeat_keep(paragraph, keep_next=False, keep_lines=False, page_break_before=False):
    fmt = paragraph.paragraph_format
    fmt.keep_with_next = keep_next
    fmt.keep_together = keep_lines
    fmt.page_break_before = page_break_before


def add_page_field(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr_text, fld_char2])
    set_run_font(run, size=8.5, color=MUTED)


def add_hyperlink(paragraph, text, url, color=BLUE, underline=True):
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), "Calibri")
    r_fonts.set(qn("w:hAnsi"), "Calibri")
    r_pr.append(r_fonts)
    color_el = OxmlElement("w:color")
    color_el.set(qn("w:val"), color)
    r_pr.append(color_el)
    if underline:
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        r_pr.append(u)
    new_run.append(r_pr)
    text_el = OxmlElement("w:t")
    text_el.text = text
    new_run.append(text_el)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


def add_numbering_definitions(doc: Document):
    numbering = doc.part.numbering_part.element
    existing_abs = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    existing_num = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    abs_bullet = max(existing_abs, default=0) + 1
    abs_decimal = abs_bullet + 1
    num_bullet = max(existing_num, default=0) + 1
    num_decimal = num_bullet + 1

    def abstract(abstract_id, num_fmt, lvl_text, font=None):
        abs_el = OxmlElement("w:abstractNum")
        abs_el.set(qn("w:abstractNumId"), str(abstract_id))
        multi = OxmlElement("w:multiLevelType")
        multi.set(qn("w:val"), "singleLevel")
        abs_el.append(multi)
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        lvl.append(start)
        fmt = OxmlElement("w:numFmt")
        fmt.set(qn("w:val"), num_fmt)
        lvl.append(fmt)
        text_el = OxmlElement("w:lvlText")
        text_el.set(qn("w:val"), lvl_text)
        lvl.append(text_el)
        jc = OxmlElement("w:lvlJc")
        jc.set(qn("w:val"), "left")
        lvl.append(jc)
        p_pr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), "720")
        tabs.append(tab)
        p_pr.append(tabs)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), "720")
        ind.set(qn("w:hanging"), "360")
        p_pr.append(ind)
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:after"), "160")
        spacing.set(qn("w:line"), "280")
        spacing.set(qn("w:lineRule"), "auto")
        p_pr.append(spacing)
        lvl.append(p_pr)
        if font:
            r_pr = OxmlElement("w:rPr")
            r_fonts = OxmlElement("w:rFonts")
            r_fonts.set(qn("w:ascii"), font)
            r_fonts.set(qn("w:hAnsi"), font)
            r_pr.append(r_fonts)
            lvl.append(r_pr)
        abs_el.append(lvl)
        numbering.append(abs_el)

    def num_instance(num_id, abstract_id):
        num = OxmlElement("w:num")
        num.set(qn("w:numId"), str(num_id))
        abs_id = OxmlElement("w:abstractNumId")
        abs_id.set(qn("w:val"), str(abstract_id))
        num.append(abs_id)
        numbering.append(num)

    abstract(abs_bullet, "bullet", "•", "Symbol")
    abstract(abs_decimal, "decimal", "%1.")
    num_instance(num_bullet, abs_bullet)
    num_instance(num_decimal, abs_decimal)
    return num_bullet, num_decimal


def apply_num(paragraph, num_id, ilvl=0):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl_el = OxmlElement("w:ilvl")
    ilvl_el.set(qn("w:val"), str(ilvl))
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl_el, num_id_el])


def set_doc_defaults(doc: Document):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    for style_name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.0
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    title = styles["Title"]
    title.font.name = "Calibri"
    title._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    title._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    title.font.size = Pt(29)
    title.font.bold = True
    title.font.color.rgb = rgb(NAVY)
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(10)
    title.paragraph_format.line_spacing = 1.0

    subtitle = styles["Subtitle"]
    subtitle.font.name = "Calibri"
    subtitle._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    subtitle._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    subtitle.font.size = Pt(14)
    subtitle.font.color.rgb = rgb(MUTED)
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after = Pt(8)

    if "Table Text" not in styles:
        table_text = styles.add_style("Table Text", WD_STYLE_TYPE.PARAGRAPH)
    else:
        table_text = styles["Table Text"]
    table_text.font.name = "Calibri"
    table_text._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    table_text._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    table_text.font.size = Pt(9)
    table_text.font.color.rgb = rgb(INK)
    table_text.paragraph_format.space_before = Pt(0)
    table_text.paragraph_format.space_after = Pt(0)
    table_text.paragraph_format.line_spacing = 1.05

    if "Source Text" not in styles:
        source = styles.add_style("Source Text", WD_STYLE_TYPE.PARAGRAPH)
    else:
        source = styles["Source Text"]
    source.font.name = "Calibri"
    source._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    source._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    source.font.size = Pt(9)
    source.font.color.rgb = rgb(MUTED)
    source.paragraph_format.space_before = Pt(4)
    source.paragraph_format.space_after = Pt(4)
    source.paragraph_format.line_spacing = 1.0

    settings = doc.settings.element
    theme_lang = settings.find(qn("w:themeFontLang"))
    if theme_lang is None:
        theme_lang = OxmlElement("w:themeFontLang")
        settings.append(theme_lang)
    theme_lang.set(qn("w:val"), "tr-TR")


def configure_section(section, header_footer=True):
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(PAGE_WIDTH_IN)
    section.page_height = Inches(PAGE_HEIGHT_IN)
    section.top_margin = Inches(MARGIN_IN)
    section.right_margin = Inches(MARGIN_IN)
    section.bottom_margin = Inches(MARGIN_IN)
    section.left_margin = Inches(MARGIN_IN)
    section.header_distance = Inches(HEADER_FOOTER_IN)
    section.footer_distance = Inches(HEADER_FOOTER_IN)
    if not header_footer:
        section.header.is_linked_to_previous = False
        section.footer.is_linked_to_previous = False
        section.header.paragraphs[0].text = ""
        section.footer.paragraphs[0].text = ""


def set_main_header_footer(section):
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False

    hp = section.header.paragraphs[0]
    hp.text = ""
    hp.paragraph_format.space_before = Pt(0)
    hp.paragraph_format.space_after = Pt(0)
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = hp.add_run("OY DASHBOARD  |  VERİ YERLEŞİMİ VE AI ALTYAPI STRATEJİSİ")
    set_run_font(run, size=8.5, color=MUTED, bold=True)

    fp = section.footer.paragraphs[0]
    fp.text = ""
    fp.paragraph_format.space_before = Pt(0)
    fp.paragraph_format.space_after = Pt(0)
    fp.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)
    r = fp.add_run("Kurum içi kullanım")
    set_run_font(r, size=8.5, color=MUTED)
    r2 = fp.add_run("\tSayfa ")
    set_run_font(r2, size=8.5, color=MUTED)
    add_page_field(fp)

    sect_pr = section._sectPr
    pg_num_type = sect_pr.find(qn("w:pgNumType"))
    if pg_num_type is None:
        pg_num_type = OxmlElement("w:pgNumType")
        sect_pr.append(pg_num_type)
    pg_num_type.set(qn("w:start"), "1")


def add_body_paragraph(doc, text, bold_lead=None, italic=False, color=INK, after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.10
    if bold_lead and text.startswith(bold_lead):
        lead = p.add_run(bold_lead)
        set_run_font(lead, bold=True, color=color)
        rest = p.add_run(text[len(bold_lead):])
        set_run_font(rest, italic=italic, color=color)
    else:
        run = p.add_run(text)
        set_run_font(run, italic=italic, color=color)
    return p


def add_bullet(doc, text, num_bullet, bold_lead=None, color=INK):
    p = doc.add_paragraph()
    apply_num(p, num_bullet)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.167
    p.paragraph_format.keep_together = True
    if bold_lead and text.startswith(bold_lead):
        r1 = p.add_run(bold_lead)
        set_run_font(r1, bold=True, color=color)
        r2 = p.add_run(text[len(bold_lead):])
        set_run_font(r2, color=color)
    else:
        r = p.add_run(text)
        set_run_font(r, color=color)
    return p


def add_numbered(doc, text, num_decimal, bold_lead=None):
    p = doc.add_paragraph()
    apply_num(p, num_decimal)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.167
    p.paragraph_format.keep_together = True
    if bold_lead and text.startswith(bold_lead):
        r1 = p.add_run(bold_lead)
        set_run_font(r1, bold=True)
        r2 = p.add_run(text[len(bold_lead):])
        set_run_font(r2)
    else:
        r = p.add_run(text)
        set_run_font(r)
    return p


def add_callout(doc, label, text, fill=CALLOUT, accent=BLUE, compact=False):
    table = doc.add_table(rows=1, cols=1)
    # Named callout override: 180 DXA inset matches the wider callout cell padding.
    set_table_geometry(table, [CONTENT_WIDTH_DXA], indent_dxa=180)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_borders(cell, color=accent, size="10")
    set_cell_margins(cell, top=140, bottom=140, start=180, end=180)
    p = cell.paragraphs[0]
    p.style = doc.styles["Normal"]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.10
    r1 = p.add_run(label.upper() + "  ")
    set_run_font(r1, size=9 if compact else 9.5, color=accent, bold=True)
    r2 = p.add_run(text)
    set_run_font(r2, size=9.5 if compact else 10.5, color=INK, bold=not compact)
    p.paragraph_format.keep_together = True
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(0)
    spacer.paragraph_format.space_before = Pt(0)
    spacer.add_run("")
    return table


def set_table_text(cell, text, bold=False, color=INK, size=9, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = cell.paragraphs[0]
    p.style = "Table Text"
    p.alignment = align
    p.paragraph_format.keep_together = True
    p.paragraph_format.keep_with_next = False
    p.text = ""
    r = p.add_run(str(text))
    set_run_font(r, size=size, color=color, bold=bold)
    return p


def add_table(doc, headers, rows, widths_dxa, header_fill=LIGHT_GRAY, font_size=9, alignments=None):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths_dxa)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = table.rows[0]
    repeat_table_header(hdr)
    prevent_row_split(hdr)
    for i, header in enumerate(headers):
        set_cell_shading(hdr.cells[i], header_fill)
        set_table_text(hdr.cells[i], header, bold=True, color=NAVY, size=font_size,
                       align=(alignments[i] if alignments else WD_ALIGN_PARAGRAPH.LEFT))
    for row_values in rows:
        row = table.add_row()
        prevent_row_split(row)
        for i, value in enumerate(row_values):
            align = alignments[i] if alignments else WD_ALIGN_PARAGRAPH.LEFT
            set_table_text(row.cells[i], value, size=font_size, align=align)
    set_table_geometry(table, widths_dxa)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    return table


def add_citation_run(paragraph, key):
    r = paragraph.add_run(f" [{key}]")
    set_run_font(r, size=7.5, color=BLUE, bold=True)
    r.font.superscript = True


def page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def draw_architecture_diagram(path: Path):
    path.parent.mkdir(parents=True, exist_ok=True)
    W, H = 1700, 1030
    img = Image.new("RGB", (W, H), "white")
    d = ImageDraw.Draw(img)
    font_dir = Path("C:/Windows/Fonts")

    def f(name, size):
        candidate = font_dir / name
        if not candidate.exists():
            candidate = font_dir / "arial.ttf"
        return ImageFont.truetype(str(candidate), size)

    title_f = f("arialbd.ttf", 40)
    region_f = f("arialbd.ttf", 28)
    box_f = f("arialbd.ttf", 23)
    body_f = f("arial.ttf", 20)
    small_f = f("arial.ttf", 17)

    def rounded_box(xy, fill, outline, title, lines=(), radius=18, title_color=NAVY):
        d.rounded_rectangle(xy, radius=radius, fill=f"#{fill}", outline=f"#{outline}", width=3)
        x1, y1, x2, y2 = xy
        d.text((x1 + 18, y1 + 16), title, font=box_f, fill=f"#{title_color}")
        y = y1 + 56
        for line in lines:
            d.text((x1 + 18, y), line, font=small_f, fill="#344054")
            y += 27

    def arrow(a, b, color=BLUE, width=5):
        d.line([a, b], fill=f"#{color}", width=width)
        bx, by = b
        ax, ay = a
        import math
        angle = math.atan2(by - ay, bx - ax)
        length = 17
        for delta in (2.55, -2.55):
            x = bx + length * math.cos(angle + delta)
            y = by + length * math.sin(angle + delta)
            d.line([b, (x, y)], fill=f"#{color}", width=width)

    d.text((60, 35), "Önerilen Türkiye İçi Veri ve AI Çalışma Mimarisi", font=title_f, fill=f"#{NAVY}")
    d.text((60, 88), "Üretim verisi, felaket kurtarma ve model eğitimi aynı veri yerleşimi politikası altında yönetilir.", font=body_f, fill="#475467")

    rounded_box((60, 175, 390, 300), LIGHT_GRAY, BORDER, "Mobil / Web Uygulaması", ("TLS, MFA, cihaz güvenliği",))
    rounded_box((500, 155, 1630, 710), "F8FBFF", BLUE, "ANKARA - BİRİNCİL ORTAM", (), title_color=BLUE)
    rounded_box((570, 235, 1030, 345), LIGHT_BLUE, BLUE, "API ve Uygulama Platformu", ("Yük dengeleme • WAF • konteyner/VM",))
    rounded_box((1090, 235, 1540, 345), CALLOUT, BORDER, "Rıza ve Denetim Kayıtları", ("Sürüm, zaman, amaç, erişim izi",))
    rounded_box((570, 405, 850, 555), PALE_GOLD, "D9B75E", "Kimlik Kasası", ("Doğrudan tanımlayıcılar", "ayrı anahtar alanı"), title_color=GOLD)
    rounded_box((890, 405, 1170, 555), PALE_GREEN, "7FB08D", "Sağlık Veritabanı", ("person_id ile", "pseudonim kayıtlar"), title_color=GREEN)
    rounded_box((1210, 405, 1540, 555), LIGHT_BLUE, BLUE, "Nesne Deposu", ("Fotoğraf, video, 3B tarama", "sürümleme ve şifreleme"))
    rounded_box((720, 600, 1390, 685), CALLOUT, BORDER, "Türkiye İçi Anahtar Yönetimi / HSM", ("Müşteri kontrollü anahtar • görev ayrılığı",))

    rounded_box((60, 430, 390, 655), PALE_GREEN, "7FB08D", "Mahremiyet ve Veri Hazırlama Hattı", ("Tanımlayıcı temizleme", "EXIF / serbest metin kontrolü", "yeniden tanımlama testi", "veri seti sürümleme"), title_color=GREEN)
    rounded_box((60, 770, 760, 965), PALE_RED, "D98B8B", "İZOLE AI EĞİTİM ALANI - TÜRKİYE", ("GPUaaS veya dedike GPU • dış ağ çıkışı varsayılan kapalı", "anonim veri varsayılanı • model kayıt sistemi • sızıntı testleri"), title_color=RED)
    rounded_box((900, 770, 1630, 965), LIGHT_GRAY, BORDER, "İSTANBUL - FELAKET KURTARMA", ("Asenkron veritabanı replikası • değiştirilemez yedek", "hedef: RPO ≤ 15 dk • RTO ≤ 4 saat • düzenli geri dönüş testi"))

    arrow((390, 238), (570, 288))
    arrow((1030, 288), (1090, 288))
    arrow((800, 345), (710, 405))
    arrow((800, 345), (1030, 405))
    arrow((1030, 345), (1370, 405))
    arrow((500, 520), (390, 540), color=GREEN)
    arrow((230, 655), (330, 770), color=GREEN)
    arrow((1170, 555), (1180, 770), color=RED)
    arrow((1370, 555), (1390, 770), color=MUTED)
    d.text((1040, 725), "Şifreli replikasyon / yedek", font=small_f, fill="#475467")
    d.text((80, 980), "Temel ilke: Kod veriye gelir; sağlık verisi kontrolsüz biçimde servis veya ülke değiştirmez.", font=small_f, fill=f"#{NAVY}")
    img.save(path, "PNG")


def add_picture_with_alt(doc, path: Path, alt_text: str, width=Inches(6.5)):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run()
    inline = run.add_picture(str(path), width=width)
    doc_pr = inline._inline.docPr
    doc_pr.set("descr", alt_text)
    return p


def build_document():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    TMP_DIR.mkdir(parents=True, exist_ok=True)
    draw_architecture_diagram(DIAGRAM_FILE)

    doc = Document()
    set_doc_defaults(doc)
    num_bullet, num_decimal = add_numbering_definitions(doc)

    props = doc.core_properties
    props.title = "OY Dashboard - Türkiye İçinde Sağlık Verisi Barındırma ve Yapay Zekâ Eğitim Altyapısı Strateji Raporu"
    props.subject = "Veri yerleşimi, KVKK uyumu, bulut sağlayıcı seçimi ve uygulama yol haritası"
    props.author = "OY Dashboard"
    props.keywords = "KVKK, sağlık verisi, Türkiye, veri merkezi, yapay zekâ, veri yönetişimi, felaket kurtarma"
    props.comments = "Kurum içi karar ve satın alma hazırlığı için strateji raporu."

    # Cover section: editorial_cover pattern, no running furniture.
    cover = doc.sections[0]
    configure_section(cover, header_footer=False)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(64)
    p.paragraph_format.space_after = Pt(18)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("STRATEJİ RAPORU")
    set_run_font(r, size=10, color=BLUE, bold=True)

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("Türkiye İçinde Sağlık Verisi\nBarındırma ve Yapay Zekâ\nEğitim Altyapısı")

    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("OY Dashboard için veri yerleşimi, güvenlik, sağlayıcı seçimi ve uygulama yol haritası")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(26)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("Karar taslağı  |  Sürüm 1.0")
    set_run_font(r, size=11, color=NAVY, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(80)
    r = p.add_run("28 Ağustos 2026")
    set_run_font(r, size=11, color=MUTED)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(96)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("HAZIRLANAN KURUM")
    set_run_font(r, size=8.5, color=MUTED, bold=True)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(2)
    r = p2.add_run("OY Dashboard")
    set_run_font(r, size=13, color=NAVY, bold=True)
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p3.add_run("Altyapı ve Veri Stratejisi")
    set_run_font(r, size=10, color=MUTED)

    # Main section starts at page 1.
    main = doc.add_section(WD_SECTION.NEW_PAGE)
    configure_section(main)
    set_main_header_footer(main)

    h = doc.add_paragraph("Yönetici Özeti", style="Heading 1")
    h.paragraph_format.space_before = Pt(0)

    add_callout(
        doc,
        "Karar önerisi",
        "Birinci aşamada Türk Telekom üzerinde Ankara birincil ortam, İstanbul felaket kurtarma ve Türkiye içi GPU hizmetiyle pilot kurulması; Turkcell Bulut ile NGN Cloud'un aynı şartname üzerinden ticari ve teknik karşı teklif olarak değerlendirilmesi önerilir.",
        fill=LIGHT_BLUE,
        accent=BLUE,
    )

    p = add_body_paragraph(doc, "OY Dashboard büyüdükçe ayak sağlığına, anatomik ölçümlere, görüntülere ve 3B taramalara ilişkin yüksek hacimli veri işleyecektir. Bu veri kümesinin önemli bölümü sağlık verisi niteliğindedir. Dolayısıyla altyapı kararı yalnızca kapasite ve maliyet üzerinden değil; hukuki işleme amacı, veri yerleşimi, şifreleme, erişim denetimi, felaket kurtarma ve yapay zekâ veri yaşam döngüsü birlikte ele alınarak verilmelidir.")
    add_citation_run(p, "K1")
    p = add_body_paragraph(doc, "Önerilen model, verinin Türkiye'de tutulmasını bir pazarlama beyanı olmaktan çıkarıp sözleşme, mimari ve operasyon kontrolleriyle doğrulanabilir bir kurala dönüştürür. Ana kopya kadar replikalar, yedekler, loglar, şifreleme anahtarları, destek erişimleri ve AI eğitim süreçleri de kapsam içindedir.")
    add_citation_run(p, "K5")

    doc.add_paragraph("Temel kararlar", style="Heading 2")
    add_bullet(doc, "Barındırma modeli: Türkiye içi yönetilen bulut; ilk aşamada fiziksel sunucu satın alma veya colocation önerilmez.", num_bullet, "Barındırma modeli:")
    add_bullet(doc, "Coğrafi süreklilik: Birincil ve felaket kurtarma ortamları farklı şehirlerde kurulmalıdır; önerilen başlangıç hedefi RPO ≤ 15 dakika ve RTO ≤ 4 saattir.", num_bullet, "Coğrafi süreklilik:")
    add_bullet(doc, "Veri ayrıştırma: Kimlik verileri, sağlık kayıtları, medya/3B dosyalar, rıza kayıtları ve AI eğitim setleri ayrı güvenlik alanlarında yönetilmelidir.", num_bullet, "Veri ayrıştırma:")
    add_bullet(doc, "AI kullanım ilkesi: Eğitimde anonim veri varsayılan olmalı; pseudonim veri yalnızca belgelenmiş hukuki dayanak, etki değerlendirmesi ve kapalı eğitim alanı ile kullanılmalıdır.", num_bullet, "AI kullanım ilkesi:")
    add_bullet(doc, "Tedarik yaklaşımı: Sağlayıcı seçimi, Türkiye içi veri yerleşimini bağlayıcı şekilde garanti edemeyen adayları baştan eleyen iki aşamalı RFP ve pilot yöntemiyle yapılmalıdır.", num_bullet, "Tedarik yaklaşımı:")

    doc.add_paragraph("Raporun Kullanım Amacı", style="Heading 2")
    add_body_paragraph(doc, "Bu rapor; yönetim kararı, teknik tasarım, KVKK uyum hazırlığı ve sağlayıcı satın alma süreci için ortak bir çerçeve sunar. Hukuki mütalaa yerine geçmez. Canlıya geçişten önce veri işleme amaçları ile açık rıza/aydınlatma metinleri KVKK alanında uzman hukuk danışmanı tarafından doğrulanmalıdır.")

    doc.add_paragraph("Rapor Yapısı", style="Heading 2")
    for item in (
        "Stratejik hedef ve tasarım ilkeleri",
        "Mevzuat ve veri yönetişimi çerçevesi",
        "Veri sınıflandırması ve hedef mimari",
        "Sağlayıcı karşılaştırması ve seçim önerisi",
        "90 günlük uygulama yol haritası",
        "Güvenlik kontrolleri, AI veri yaşam döngüsü ve tedarik şartları",
        "Riskler, karar kapıları, tanımlar ve kaynaklar",
    ):
        add_numbered(doc, item, num_decimal)

    page_break(doc)

    doc.add_paragraph("1. Stratejik Hedef ve Tasarım İlkeleri", style="Heading 1")
    add_body_paragraph(doc, "Altyapı hedefi, yalnızca veriyi Türkiye sınırları içinde saklamak değildir. Esas hedef; verinin yaşam döngüsünün tamamında hangi amaçla, kim tarafından, nerede ve ne kadar süreyle işlendiğinin kanıtlanabildiği; hizmet kesintilerine dayanıklı; sağlayıcı bağımlılığı yönetilmiş bir çalışma modeli kurmaktır.")

    doc.add_paragraph("1.1 Hedeflenen sonuçlar", style="Heading 2")
    for text in (
        "Sağlık ve anatomik verilerin Türkiye'deki veri merkezlerinde tutulması ve işlenmesi.",
        "Bir veri merkezi veya şehir devre dışı kaldığında hizmetin tanımlı süre içinde geri dönebilmesi.",
        "Kimlik ile sağlık verisi arasındaki bağın minimum yetkili bileşen dışında görünmemesi.",
        "AI eğitiminin üretim sisteminden kontrollü, izlenebilir ve geri alınabilir veri setleriyle beslenmesi.",
        "Saklama süresi dolan veya işleme şartı ortadan kalkan verinin yedekler dahil yönetilebilir biçimde silinmesi.",
        "Sağlayıcı değişikliği halinde verinin standart formatlarda ve makul maliyetle dışarı alınabilmesi.",
    ):
        add_bullet(doc, text, num_bullet)

    doc.add_paragraph("1.2 Tasarım ilkeleri", style="Heading 2")
    principles = [
        ("Türkiye içi varsayılan", "Yeni servis ve veri akışı, Türkiye dışına çıkamayacak şekilde varsayılan kapalı tasarlanır."),
        ("Asgari veri", "Her ekran, API ve model yalnızca ihtiyacı olan veri alanına erişir."),
        ("Kimlik ayrıştırma", "Doğrudan tanımlayıcılar klinik/anatomik kayıtlardan ayrı tutulur."),
        ("Müşteri kontrollü anahtar", "Şifreleme anahtarları mümkünse sağlayıcının genel yönetim alanından ayrılır ve Türkiye'de tutulur."),
        ("Geri yükleme kanıtı", "Yedek alınması değil, yedeğin düzenli ve ölçümlü biçimde geri döndürülebilmesi başarı ölçütüdür."),
        ("Kod veriye gelir", "Eğitim kodu kapalı veri alanına taşınır; ham sağlık verisi dış servislere taşınmaz."),
        ("Çıkış planı", "Veri formatı, anahtar devri, silme belgesi ve göç süresi daha sözleşme aşamasında tanımlanır."),
    ]
    add_table(doc, ["İlke", "Uygulama karşılığı"], principles, [2700, 6660], header_fill=LIGHT_BLUE, font_size=9.5)

    doc.add_paragraph("2. Mevzuat ve Veri Yönetişimi Çerçevesi", style="Heading 1")
    p = add_body_paragraph(doc, "Ayak sağlığı değerlendirmeleri, klinik şikâyetler, basınç/denge ölçümleri ve anatomik bulgular kişinin sağlığına ilişkin bilgi üretiyorsa özel nitelikli kişisel veri olarak ele alınmalıdır. Fotoğraf veya 3B tarama bir kişiyi benzersiz biçimde tanımak için kullanılıyorsa biyometrik veri değerlendirmesi de gündeme gelebilir.")
    add_citation_run(p, "K1")

    doc.add_paragraph("2.1 İşleme şartı ve amaç ayrımı", style="Heading 2")
    add_body_paragraph(doc, "Üretim hizmeti için toplanan sağlık verisinin model geliştirme, araştırma veya ürün optimizasyonu amacıyla kullanılması ayrı bir işleme amacı olarak değerlendirilmelidir. Sağlık hizmetinin yürütülmesine dayanan hukuki gerekçe, AR-GE veya AI eğitimi amacını kendiliğinden kapsamaz. Uygun başka bir işleme şartı bulunmadığı durumda AI eğitimi için ayrı, belirli ve bilgilendirmeye dayalı açık rıza en ihtiyatlı uygulamadır. Rızanın hizmetin zorunlu koşulu gibi sunulmaması ve geri çekme sonucunun veri seti/model yaşam döngüsüne yansıtılabilmesi gerekir.")
    add_callout(doc, "Kritik ayrım", "Pseudonimleştirme KVKK kapsamını ortadan kaldırmaz. Veri, ek bilgi veya eşleştirme yoluyla kişiye yeniden bağlanabiliyorsa kişisel veri olmaya devam eder.", fill=PALE_GOLD, accent=GOLD, compact=True)

    doc.add_paragraph("2.2 Zorunlu yönetişim çıktıları", style="Heading 2")
    governance_rows = [
        ("Kişisel veri işleme envanteri", "Veri alanı, kaynak, amaç, hukuki dayanak, alıcı, sistem, saklama süresi ve güvenlik tedbiri."),
        ("Aydınlatma ve rıza kayıtları", "Metin sürümü, zaman, kanal, kapsam, geri çekme ve ispat kaydı."),
        ("Saklama ve imha politikası", "Üretim, yedek, log, veri seti ve model çıktıları için ayrı süre/işlem."),
        ("Mahremiyet etki değerlendirmesi", "Yüksek riskli sağlık ve AI süreçlerinde risk, gereklilik, orantılılık ve azaltıcı kontrol."),
        ("Veri işleyen sözleşmeleri", "Bulut, destek, güvenlik, analitik ve AI alt yüklenicilerinin rol ve sınırları."),
        ("İhlal müdahale planı", "Tespit, sınıflandırma, delil koruma, yönetim/hukuk eskalasyonu ve 72 saatlik bildirim takvimi."),
    ]
    add_table(doc, ["Çıktı", "Asgari kapsam"], governance_rows, [2850, 6510], header_fill=LIGHT_GRAY, font_size=9.2)
    p = add_body_paragraph(doc, "Bulut sağlayıcısı veri işleyen konumunda olsa da güvenlik ve denetim sorumluluğu tamamen sağlayıcıya devredilemez. Veri sorumlusu, alınan teknik ve idari tedbirlerin yeterliliğini değerlendirmek ve izlemekle yükümlüdür.")
    add_citation_run(p, "K6")

    doc.add_paragraph("2.3 VERBİS ve kayıt eşiği", style="Heading 2")
    p = add_body_paragraph(doc, "2025 tarihli güncel istisna yaklaşımında ana faaliyet konusu özel nitelikli kişisel veri işleme olan işletmeler bakımından çalışan sayısı 10'dan az ve yıllık mali bilanço toplamı 10 milyon TL'den az ise Sicil istisnası gündeme gelebilir. Bu eşiklerden birinin aşılması veya faaliyet yapısının değişmesi halinde kayıt yükümlülüğü yeniden değerlendirilmelidir.")
    add_citation_run(p, "K7")

    page_break(doc)

    doc.add_paragraph("3. Veri Sınıflandırması ve Yaşam Döngüsü", style="Heading 1")
    add_body_paragraph(doc, "Veri tabanı tasarımından önce sınıflandırma yapılmalıdır. Aynı güvenlik seviyesindeki verileri tek yerde toplamak yerine, işlev ve risk düzeyine göre ayrıştırmak hem ihlal etkisini hem de gereksiz erişimi azaltır.")

    data_rows = [
        ("D0 - Açık", "Genel ürün metinleri, anonim istatistik", "Standart", "Kurumsal yayın politikası"),
        ("D1 - İç", "Sistem konfigürasyonu, anonim operasyon metriği", "İç kullanıcı", "İş ihtiyacı kadar"),
        ("D2 - Kişisel", "Ad, e-posta, telefon, hesap bilgisi", "Sınırlı", "Amaç/mevzuat süresi"),
        ("D3 - Özel nitelikli", "Ayak sağlığı, tanı/şikâyet, görüntü, 3B tarama", "Yüksek", "Kısa ve gerekçeli"),
        ("D4 - Kritik güvenlik", "Anahtar, erişim sırrı, kimlik eşleştirme tablosu", "En yüksek", "Rotasyon ve görev ayrılığı"),
    ]
    add_table(doc, ["Sınıf", "Örnek", "Koruma", "Saklama yaklaşımı"], data_rows, [1450, 3500, 1450, 2960], header_fill=LIGHT_BLUE, font_size=8.7)

    doc.add_paragraph("3.1 Önerilen mantıksal veri ayrımı", style="Heading 2")
    components = (
        ("Kimlik Kasası", "Ad, iletişim, kimlik ve hesap eşleştirme bilgilerinin tutulduğu ayrı veri alanı. Sağlık kaydına yalnızca rastgele üretilmiş person_id ile bağlanır."),
        ("Sağlık Veritabanı", "Yapısal ölçüm, şikâyet, değerlendirme ve klinik sonuçların bulunduğu yüksek güvenlikli ilişkisel veri alanı."),
        ("Nesne Deposu", "Fotoğraf, video ve 3B taramaların tutulduğu sürümlü, şifreli ve yaşam döngüsü kurallı depolama."),
        ("Rıza ve Denetim Defteri", "Aydınlatma/rıza sürümleri ile okuma, değiştirme, indirme ve yönetici erişimlerinin değiştirilemez kaydı."),
        ("AI Veri Gölü", "Üretimden doğrudan sorgulanmayan; kontrollü dışa aktarım hattıyla oluşan, sürümlenmiş ve kalite kayıtlı eğitim setleri."),
        ("Model Kayıt Sistemi", "Veri seti sürümü, kod sürümü, parametre, metrik, onay, risk testi ve yayına alma kararının tutulduğu alan."),
    )
    for name, desc in components:
        add_bullet(doc, f"{name}: {desc}", num_bullet, f"{name}:")

    doc.add_paragraph("3.2 Saklama ve imha", style="Heading 2")
    add_body_paragraph(doc, "Tek bir 'tüm veriler için X yıl' süresi belirlenmemelidir. Hesap, sağlık kaydı, görüntü, log, rıza ispatı, yedek ve AI veri seti için ayrı saklama kriteri yazılmalıdır. Süre dolduğunda üretim kaydı silinir; replikalar, yedekler ve veri setleri için de erişilemez hâle getirme ve sonrasında kesin imha takvimi işletilir. Anonim istatistikler, yeniden tanımlama riski periyodik olarak test edilmek şartıyla daha uzun süre tutulabilir.")

    doc.add_paragraph("4. Hedef Teknik Mimari", style="Heading 1")
    add_picture_with_alt(doc, DIAGRAM_FILE, "Ankara birincil ortam, İstanbul felaket kurtarma ve Türkiye içi izole yapay zekâ eğitim alanını gösteren hedef mimari diyagramı")
    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(0)
    caption.paragraph_format.space_after = Pt(8)
    r = caption.add_run("Şekil 1. Türkiye içi üretim, felaket kurtarma ve AI eğitim mimarisi")
    set_run_font(r, size=9, color=MUTED, italic=True)

    doc.add_paragraph("4.1 Birincil üretim ortamı", style="Heading 2")
    add_body_paragraph(doc, "Varsayılan yerleşim Ankara olarak önerilir. Uygulama sunucuları veya Kubernetes kümesi, WAF/yük dengeleme, yüksek erişilebilir veritabanı, nesne deposu ve yerel anahtar yönetimi aynı güvenlik sınırı içinde fakat ayrı ağ segmentlerinde çalışır. Yönetim erişimi internete açık olmamalı; kurumsal VPN, MFA ve ayrıcalıklı erişim yönetimi üzerinden sağlanmalıdır.")

    doc.add_paragraph("4.2 Felaket kurtarma", style="Heading 2")
    add_body_paragraph(doc, "İstanbul ortamında asenkron veritabanı replikası, değiştirilemez yedek ve asgari uygulama kapasitesi tutulur. Aynı şehir veya kampüsteki ikinci kabin felaket kurtarma sayılmaz. Geçiş işlemi belgelenmeli; yılda en az iki kez tam senaryo, aylık olarak da örnek geri yükleme testi yapılmalıdır.")

    doc.add_paragraph("4.3 Log, izleme ve destek verisi", style="Heading 2")
    add_body_paragraph(doc, "Uygulama loglarında sağlık bilgisi, fotoğraf yolu, serbest metin veya doğrudan tanımlayıcı bulunmamalıdır. Hata izleme araçlarına yalnızca anonim teknik bağlam gönderilmelidir. SIEM, APM ve destek kaydı sağlayıcıları da veri yerleşimi envanterine dahil edilir; destek ekran görüntüleri ve ticket ekleri ayrı politika ile sınırlandırılır.")

    page_break(doc)

    doc.add_paragraph("5. Türkiye'deki Sağlayıcıların Ön Değerlendirmesi", style="Heading 1")
    add_body_paragraph(doc, "Aşağıdaki değerlendirme 28 Ağustos 2026 itibarıyla kamuya açık teknik beyanlara dayanır. Fiyat, gerçek hizmet seviyesi, alt yüklenici yapısı ve veri yerleşimi ancak sözleşme ve pilotla doğrulanabilir.")

    provider_rows = [
        ("Türk Telekom", "Birincil aday", "İstanbul-Ankara veri merkezi; replikasyon, yönetilen hizmet, WAF/DDoS ve Türkiye içi GPU hizmeti", "GPU modeli/kapasitesi, müşteri kontrollü anahtar ve her servisin kesin lokasyonu teyit edilmeli", "Pilot"),
        ("Turkcell Bulut", "Güçlü alternatif", "Geniş veri merkezi ağı; nesne depolama, yönetilen veritabanı, Kubernetes ve iş sürekliliği", "Türkiye içi GPU eğitim hizmetinin modeli ve ticari koşulları teyit edilmeli", "RFP"),
        ("NGN Cloud", "Üçüncü aday", "Operatör bağımsız Tier III merkez; İstanbul ve yeni Ankara bölgesi; yerli AI Compute yatırımı", "Yeni bölge/AI hizmetinin olgunluğu, kapasite ve SLA kanıtı istenmeli", "RFP"),
        ("AWS İstanbul LZ", "Yardımcı iş yükü", "Yerel EC2/EBS/S3/EKS/ECS; AWS araçlarıyla uyum", "Tam Region değil; Frankfurt ana bölgesi; varsayılan snapshot davranışı ve servis kapsamı dikkat gerektirir", "Sınırlı"),
        ("Radore", "Geliştirme/ikincil", "Kolay ve görece ekonomik bulut sunucu seçenekleri", "Kamuya açık ana merkez tek İstanbul lokasyonu; bağımsız DR olmadan ana sağlık sistemi için yetersiz", "Duruma bağlı"),
        ("Equinix", "Colocation", "Operatör bağımsız, yüksek bağlantı kapasitesi ve kurumsal standartlar", "Donanım, işletim sistemi, yedek ve güvenlik operasyonu büyük ölçüde müşteri sorumluluğunda", "İleri ölçek"),
    ]
    add_table(doc, ["Sağlayıcı", "Rol", "Güçlü yön", "Doğrulanacak konu", "Sonuç"], provider_rows,
              [1400, 1300, 2650, 3100, 910], header_fill=LIGHT_BLUE, font_size=7.9)
    src = doc.add_paragraph(style="Source Text")
    src.add_run("Kaynak çerçevesi: Türk Telekom [K8-K9], Turkcell [K10-K11], NGN [K12], AWS [K13-K14], Google Cloud/Turkcell gelecek bölge bilgisi [K15-K16].")

    doc.add_paragraph("5.1 Ağırlıklı seçim modeli", style="Heading 2")
    scoring_rows = [
        ("Türkiye içi veri yerleşimi ve sözleşme", "%25", "23", "23", "22", "Geçiş kriteri"),
        ("Coğrafi süreklilik ve DR", "%15", "14", "15", "13", "İki şehir kanıtı"),
        ("Yönetilen platform ve veri servisleri", "%15", "12", "14", "10", "PaaS kapsamı"),
        ("AI/GPU kapasitesi", "%15", "14", "9", "11", "Model ve fiyat"),
        ("Güvenlik ve denetim kanıtı", "%15", "13", "14", "13", "Rapor/SLA"),
        ("Operasyon ve destek", "%10", "8", "8", "7", "7/24 yetkinlik"),
        ("Toplam maliyet ve çıkış", "%5", "3", "3", "3", "Teklif sonrası"),
        ("Ön toplam", "%100", "87", "86", "79", "Bağlayıcı değil"),
    ]
    add_table(doc, ["Kriter", "Ağırlık", "TT", "Turkcell", "NGN", "Not"], scoring_rows,
              [3200, 900, 780, 780, 780, 2920], header_fill=LIGHT_GRAY, font_size=8.3,
              alignments=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT])
    add_callout(doc, "Yorum", "Türk Telekom ile Turkcell arasındaki fark karar vermek için yeterli değildir. Türk Telekom'un öne çıkmasının nedeni Türkiye içi GPU hizmetini bugün açıkça sunmasıdır. Turkcell, yönetilen platform kapsamı ve uzun vadeli kapasitede daha güçlü sonuç verebilir.", fill=PALE_GREEN, accent=GREEN, compact=True)

    doc.add_paragraph("6. Önerilen Sağlayıcı ve Kurulum Modeli", style="Heading 1")
    add_body_paragraph(doc, "İlk satın alma ve pilot adayı Türk Telekom olmalıdır. Bunun nedeni tek sözleşme altında iki şehir, replikasyon, yönetilen operasyon, güvenlik hizmetleri ve GPU kapasitesinin birlikte değerlendirilebilmesidir. Pilot başarı kriterlerini karşılamaz veya sözleşme tüm veri türlerinin Türkiye'de kalacağını açıkça garanti etmezse bir sonraki aday Turkcell Bulut olmalıdır.")

    doc.add_paragraph("6.1 Önerilen başlangıç kapsamı", style="Heading 2")
    for text in (
        "Ankara'da üç uygulama düğümü veya küçük bir yönetilen Kubernetes kümesi.",
        "Aynı lokasyonda senkron yüksek erişilebilir PostgreSQL; İstanbul'a asenkron replika.",
        "Fotoğraf ve 3B taramalar için sürümlü, şifreli nesne depolama.",
        "İstanbul'da değiştirilemez yedek ve minimum çalışabilir uygulama kapasitesi.",
        "WAF, DDoS koruması, merkezi log/SIEM ve 7/24 alarm yönetimi.",
        "AI eğitimi için süreli GPU kullanımı; üretim ağına doğrudan erişmeyen ayrı proje/tenant.",
        "Müşteri kontrollü anahtar veya Türkiye'deki HSM üzerinde ayrılmış anahtar alanı.",
    ):
        add_bullet(doc, text, num_bullet)

    doc.add_paragraph("6.2 Sözleşmesel geçiş kriterleri", style="Heading 2")
    for text in (
        "Birincil veri, replika, snapshot, yedek, log ve anahtarların kesin veri merkezi/lokasyon listesi sözleşme ekinde yer almalıdır.",
        "Yurt dışındaki personelin ayrıcalıklı erişimi varsayılan olarak kapalı olmalı; istisnalar ön onay ve ayrıntılı kayıt gerektirmelidir.",
        "Alt yükleniciler önceden açıklanmalı; yeni alt yüklenici veya veri lokasyonu değişikliği müşteri onayına tabi olmalıdır.",
        "İhlal veya şüpheli erişim OY Dashboard'a en geç 6-12 saat içinde bildirilmelidir.",
        "Sözleşme sonunda makinece okunabilir veri dışa aktarımı, anahtar devri ve doğrulanabilir silme belgesi sağlanmalıdır.",
    ):
        add_bullet(doc, text, num_bullet)

    page_break(doc)

    doc.add_paragraph("7. Uygulama Yol Haritası", style="Heading 1")
    add_body_paragraph(doc, "Yol haritası, gerçek sağlık verisi canlı sisteme alınmadan önce yönetişim ve güvenlik kapılarının tamamlanacağı şekilde planlanmıştır. Takvim, mevcut uygulama mimarisinin üretime hazır olduğu varsayımıyla hazırlanmıştır.")

    roadmap_rows = [
        ("0", "Hafta 0-2", "Veri envanteri, hukuki dayanak, rıza ayrımı, sorumlu rolleri", "Veri sahibi + hukuk", "Onaylı envanter ve işleme matrisi"),
        ("1", "Hafta 2-4", "Üç sağlayıcıya aynı RFP; mimari ve maliyet teklifleri", "Teknik lider + satın alma", "Kısa liste ve pilot sözleşmesi"),
        ("2", "Hafta 5-8", "Ankara/İstanbul pilotu, yedek, failover, güvenlik ve GPU testi", "Altyapı + güvenlik", "Pilot kabul raporu"),
        ("3", "Hafta 9-12", "Üretim kurulumu, sızma testi, olay tatbikatı ve kontrollü geçiş", "Teknik lider + DPO", "Canlıya geçiş onayı"),
        ("4", "Ay 4-6", "AI veri hazırlama hattı, veri seti sürümleme, model kayıt sistemi", "ML lideri + mahremiyet", "İlk kontrollü eğitim"),
        ("5", "Sürekli", "Erişim incelemesi, geri yükleme testi, kapasite ve tedarikçi denetimi", "Operasyon + yönetim", "Aylık/çeyreklik kanıt paketi"),
    ]
    add_table(doc, ["Faz", "Takvim", "Ana faaliyet", "Sorumlu", "Çıkış kriteri"], roadmap_rows,
              [650, 1100, 3250, 1850, 2510], header_fill=LIGHT_BLUE, font_size=8.2)

    doc.add_paragraph("7.1 Faz 0 - Yönetişim temeli", style="Heading 2")
    for text in (
        "Toplanan her veri alanını ve kaynağını çıkarın; zorunlu olmayan alanları kaldırın.",
        "Hizmet sunumu ile AI eğitimi amaçlarını ve hukuki dayanaklarını ayırın.",
        "Çocuk verisi ihtimali varsa veli/onay, yaş doğrulama ve ek koruma akışını tasarlayın.",
        "Veri sahibi, sistem sahibi, güvenlik sorumlusu ve model onay merciini atayın.",
        "Saklama süreleri için karar kaydı ve silme yöntemi belirleyin.",
    ):
        add_bullet(doc, text, num_bullet)

    doc.add_paragraph("7.2 Faz 1-2 - RFP ve pilot", style="Heading 2")
    for text in (
        "Aynı teknik kapasiteyi Türk Telekom, Turkcell ve NGN'ye fiyatlatın.",
        "Pilot verisini sentetik veya geri döndürülemez biçimde anonimleştirilmiş veriyle oluşturun.",
        "Yük testi, SQL yedekten dönüş, nesne geri yükleme, şehirler arası failover ve anahtar erişim testini çalıştırın.",
        "Yurt dışı ağ çıkışlarını teknik olarak engelleyip kontrollü kaçak veri testleri yapın.",
        "GPU kullanımında veri kopyalarının, geçici disklerin ve job loglarının nerede kaldığını doğrulayın.",
    ):
        add_bullet(doc, text, num_bullet)

    doc.add_paragraph("7.3 Faz 3-4 - Üretim ve AI", style="Heading 2")
    add_body_paragraph(doc, "Canlıya geçiş, kritik bulgu içermeyen sızma testi ve başarılı geri dönüş tatbikatına bağlanmalıdır. AI hattı üretimden sonra ayrı bir değişiklik olarak devreye alınmalı; veri seti oluşturma, onay, eğitim ve model yayını birbirinden ayrılmış yetkilerle yürütülmelidir.")

    doc.add_paragraph("8. Güvenlik Kontrol Çerçevesi", style="Heading 1")
    security_rows = [
        ("Kimlik ve erişim", "MFA, RBAC/ABAC, görev ayrılığı, süreli ayrıcalık, üç aylık yetki incelemesi", "IAM raporu, onay kaydı"),
        ("Şifreleme", "Aktarımda TLS; saklamada güçlü şifreleme; ayrı anahtar alanı ve rotasyon", "KMS/HSM konfigürasyonu"),
        ("Ağ güvenliği", "Özel ağ, egress allowlist, WAF, DDoS, yönetim erişiminde VPN/PAM", "Ağ diyagramı ve kural çıktısı"),
        ("Uygulama güvenliği", "SAST/DAST, bağımlılık tarama, gizli bilgi tarama, yıllık sızma testi", "CI/CD ve test raporu"),
        ("Veri güvenliği", "Sınıflandırma, DLP, log maskeleme, toplu indirme alarmı", "DLP politikası ve alarmlar"),
        ("Yedek/DR", "3-2-1-1-0 yaklaşımı, immutable kopya, aylık restore, altı aylık failover", "Test tutanağı ve süre ölçümü"),
        ("İzleme", "SIEM, değiştirilemez audit log, yönetici erişim alarmı, 7/24 olay yönetimi", "Aylık güvenlik özeti"),
        ("Tedarikçi", "Alt işleyen listesi, denetim hakkı, lokasyon değişikliği kontrolü", "Sözleşme eki ve raporlar"),
    ]
    add_table(doc, ["Kontrol alanı", "Asgari kontrol", "Kanıt"], security_rows, [1800, 5100, 2460], header_fill=LIGHT_GRAY, font_size=8.5)
    p = add_body_paragraph(doc, "KVKK Kurulunun özel nitelikli veriler için belirlediği önlemler arasında şifreleme, güvenli loglama, düzenli güvenlik testleri, yetki kontrolü ve uzaktan erişimde en az iki kademeli doğrulama bulunmaktadır.")
    add_citation_run(p, "K2")

    doc.add_paragraph("8.1 Olay müdahalesi", style="Heading 2")
    add_body_paragraph(doc, "Sağlayıcıdan gelen ilk ihbar ile Kurula yapılacak bildirimin aynı şey olmadığı unutulmamalıdır. OY Dashboard'un 72 saatlik değerlendirme ve bildirim süresini yönetebilmesi için sağlayıcı bildirim hedefi 6-12 saat olmalı; kanıtların kademeli tamamlanması beklenirken ilk bildirim geciktirilmemelidir.")
    p = doc.paragraphs[-1]
    add_citation_run(p, "K4")

    page_break(doc)

    doc.add_paragraph("9. Yapay Zekâ Eğitim Verisi Yaşam Döngüsü", style="Heading 1")
    add_body_paragraph(doc, "AI eğitim ortamı, üretim veritabanının bir kopyası olarak kurulamaz. Her veri seti belirli bir kullanım amacı, kapsam, hukuki dayanak, sürüm, sorumlu ve silme tarihiyle oluşturulmalıdır. Model eğitimi de kişisel veri işleme faaliyeti olduğundan tasarım aşamasında mahremiyet etki değerlendirmesi yapılmalıdır.")
    p = doc.paragraphs[-1]
    add_citation_run(p, "K3")

    lifecycle = [
        ("1", "Talep", "Model amacı, başarı metriği, gerekli veri alanı ve hassasiyet sınıfı yazılır."),
        ("2", "Hukuki/etik onay", "İşleme şartı, rıza kapsamı, gereklilik ve orantılılık değerlendirilir."),
        ("3", "Veri hazırlama", "Kimlik, dosya adı, EXIF, serbest metin ve nadir kombinasyonlar temizlenir."),
        ("4", "Anonimlik testi", "K-anonimlik/l-çeşitlilik/t-yakınlık gibi yöntemler veri bağlamına göre uygulanır; yeniden tanımlama riski test edilir."),
        ("5", "Eğitim", "Egress kapalı Türkiye içi GPU alanında süreli iş; geçici disk ve log yaşam döngüsü tanımlı."),
        ("6", "Model güvenliği", "Ezberleme, üyelik çıkarımı, veri sızıntısı, ayrımcılık ve performans testleri."),
        ("7", "Yayın/onay", "Model kartı, veri seti kartı, sorumlu onayı ve geri alma planı."),
        ("8", "İzleme/imha", "Drift, hata, rıza geri çekme etkisi, veri seti ve eski model sürümlerinin imhası."),
    ]
    add_table(doc, ["Adım", "Aşama", "Zorunlu çıktı"], lifecycle, [650, 1800, 6910], header_fill=LIGHT_BLUE, font_size=8.7)

    doc.add_paragraph("9.1 Anonimleştirme kontrol listesi", style="Heading 2")
    for text in (
        "Ad, iletişim, kullanıcı adı, hasta/protokol numarası ve bağlantı anahtarı çıkarıldı mı?",
        "Fotoğraf ve videolardaki yüz, çevre, belge, ekran ve diğer tanımlayıcılar temizlendi mi?",
        "Dosya adı, EXIF, GPS, cihaz kimliği ve yükleme zamanı gibi metadata kaldırıldı mı?",
        "Yaş, lokasyon, nadir hastalık ve özgün anatomik kombinasyonlar yeniden tanımlama riski yaratıyor mu?",
        "Serbest metinlerde kişi, kurum, hekim ve yer adları otomatik ve örneklemeli manuel kontrolden geçti mi?",
        "Veri seti dış kaynaklarla eşleştirildiğinde kişi belirlenebilir mi?",
        "Eğitim sonrasında modelin ham örnekleri ezberleyip üretmediği test edildi mi?",
    ):
        add_bullet(doc, text, num_bullet)

    doc.add_paragraph("10. Kapasite ve Maliyet Planlama", style="Heading 1")
    add_body_paragraph(doc, "Yerel sağlayıcıların kurumsal fiyatları çoğunlukla teklif üzerinden verildiği için karşılaştırma aylık sunucu bedeliyle sınırlanmamalıdır. Üç yıllık toplam maliyet; işlem, depolama, replikasyon, yedek, internet/egress, lisans, güvenlik, yönetilen hizmet, destek ve GPU kalemlerini içermelidir.")

    capacity_rows = [
        ("Başlangıç", "2-3 uygulama düğümü; HA veritabanı; 1 TB'a kadar nesne", "GPU yalnızca eğitim saatlerinde", "Kullanım ve veri büyümesi ölçülür"),
        ("Büyüme", "3-6 otomatik ölçeklenen düğüm; okuma replikası; 1-10 TB nesne", "Aylık paket veya ayrılmış GPU saatleri", "DR ve güvenlik otomasyonu artırılır"),
        ("Ölçek", "Kubernetes; bölümleme/analitik katman; 10 TB+ nesne", "Dedike veya çoklu GPU", "Çift sağlayıcı çıkış kopyası değerlendirilir"),
    ]
    add_table(doc, ["Aşama", "Üretim kapasitesi", "AI yaklaşımı", "Yönetim kararı"], capacity_rows,
              [1300, 3100, 2350, 2610], header_fill=LIGHT_GRAY, font_size=8.7)

    doc.add_paragraph("10.1 Kapasite formülü", style="Heading 2")
    add_callout(doc, "Planlama formülü", "Aylık ham büyüme = yeni kayıt sayısı × ortalama yapısal veri ve medya boyutu. Toplam depolama bütçesi; birincil veri + replika + sürümleme + değiştirilemez yedek + AI veri seti kopyaları ayrı ayrı hesaplanarak oluşturulmalıdır.", fill=CALLOUT, accent=BLUE, compact=True)
    add_body_paragraph(doc, "Teklifte en az üç senaryo istenmelidir: mevcut kullanım, 12 aylık tahmin ve üç yıllık yüksek büyüme. Dövizli kalemler, fiyat artış tavanı ve veri dışa aktarım bedeli ayrıca gösterilmelidir. GPU için kart modeli, VRAM, çoklu GPU ağı, minimum kiralama süresi ve boşta bekleme ücretleri ayrı satır olmalıdır.")

    page_break(doc)

    doc.add_paragraph("11. Sağlayıcı RFP ve Sözleşme Kontrol Listesi", style="Heading 1")
    add_body_paragraph(doc, "Aşağıdaki sorular her sağlayıcıya aynı biçimde yöneltilmelidir. Sözlü veya sunum içindeki cevaplar yeterli kabul edilmemeli; sözleşme eki, mimari çizim, sertifika kapsamı veya test raporuyla kanıt istenmelidir.")

    rfp_rows = [
        ("R-01", "Birincil veri, replika, snapshot, yedek, log, anahtar ve geçici kopyaların fiziksel ülkesi/lokasyonu nedir?", "Lokasyon eki"),
        ("R-02", "Türkiye dışındaki personel veya alt yüklenici ayrıcalıklı erişim sağlayabilir mi? Hangi istisnalar vardır?", "Erişim politikası"),
        ("R-03", "Veri yerleşimi veya alt yüklenici değişikliği müşteri onayına tabi midir?", "Sözleşme maddesi"),
        ("R-04", "Müşteri kontrollü anahtar, BYOK veya Türkiye içi HSM seçenekleri nelerdir?", "KMS/HSM tasarımı"),
        ("R-05", "İstanbul-Ankara arasında hangi replikasyon modeli ve ölçülebilir RPO/RTO sunulur?", "SLA ve test"),
        ("R-06", "Yedekler değiştirilemez/WORM olabilir mi; silme ve fidye yazılımı koruması nasıl uygulanır?", "Ürün dokümanı"),
        ("R-07", "PostgreSQL yönetimi, HA, sürüm güncelleme, PITR ve performans sınırları nelerdir?", "Servis kataloğu"),
        ("R-08", "Nesne depolama S3 uyumlu mu; sürümleme, yaşam döngüsü ve toplu dışa aktarım destekliyor mu?", "Teknik doğrulama"),
        ("R-09", "Kubernetes/VM altyapısında özel ağ, egress kontrolü, WAF ve DDoS hangi SLA ile sunulur?", "Mimari + SLA"),
        ("R-10", "GPU modelleri, VRAM, çoklu GPU/RDMA, süreli kullanım ve kuyruk kapasitesi nedir?", "Kapasite/fiyat"),
        ("R-11", "GPU işi sona erdiğinde geçici disk, cache, log ve snapshot nasıl ve ne zaman silinir?", "İmha prosedürü"),
        ("R-12", "ISO 27001/27017/27018/27701/22301, SOC 2 ve Tier sertifikalarının kapsamı hangi hizmetleri içerir?", "Geçerli rapor"),
        ("R-13", "Müşterinin bağımsız denetim, sızma testi veya kontrol kanıtı inceleme hakkı var mı?", "Denetim maddesi"),
        ("R-14", "Güvenlik olayı müşteriye kaç saat içinde ve hangi asgari içerikle bildirilir?", "İhlal SLA'sı"),
        ("R-15", "Sözleşme sonunda veri hangi formatta, ne kadar sürede ve hangi bedelle çıkarılır?", "Çıkış planı"),
        ("R-16", "Silme sonrası doğrulanabilir silme belgesi ve anahtar imha kanıtı sağlanır mı?", "İmha belgesi"),
        ("R-17", "Üç yıllık TCO içinde hangi kullanım, lisans, destek, egress ve artış varsayımları vardır?", "Maliyet modeli"),
    ]
    add_table(doc, ["Kod", "Gereksinim / soru", "Beklenen kanıt"], rfp_rows, [750, 6650, 1960], header_fill=LIGHT_BLUE, font_size=8.3)

    doc.add_paragraph("12. Riskler ve Karar Kapıları", style="Heading 1")
    risk_rows = [
        ("Gizli yurt dışı aktarım", "Yüksek", "Log, destek, snapshot ve analitik dahil uçtan uca veri akışı envanteri; egress kontrolü", "Pilot öncesi"),
        ("Rıza amacının yetersizliği", "Yüksek", "Hizmet ve AI amaçlarını ayırma; hukuk onayı; sürümlü rıza", "Veri toplamadan önce"),
        ("Tek sağlayıcı bağımlılığı", "Orta", "Standart veri formatı, düzenli dışa aktarım testi, çıkış maddesi", "Sözleşme öncesi"),
        ("Yedek var fakat dönmüyor", "Yüksek", "Aylık restore ve altı aylık şehir failover tatbikatı", "Canlı öncesi/sürekli"),
        ("AI veri sızıntısı", "Yüksek", "Anonimlik testi, kapalı GPU ağı, model ezberleme/çıkarım testi", "Her model öncesi"),
        ("Aşırı maliyet", "Orta", "12/36 aylık senaryo, bütçe alarmı, GPU süreli kullanım, artış tavanı", "Teklif ve aylık"),
        ("Operasyon yetkinliği", "Orta", "Yönetilen hizmet, görev tanımı, olay tatbikatı ve eğitim", "Canlı öncesi"),
    ]
    add_table(doc, ["Risk", "Seviye", "Temel azaltım", "Karar kapısı"], risk_rows,
              [2100, 1100, 4500, 1660], header_fill=LIGHT_GRAY, font_size=8.5,
              alignments=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT])

    doc.add_paragraph("12.1 Yönetim tarafından onaylanması gereken kararlar", style="Heading 2")
    for text in (
        "Veri yerleşimi hedefi: yalnızca saklama mı, yoksa destek erişimi ve işlem dahil tam Türkiye içi model mi?",
        "Önerilen iş sürekliliği hedefleri: RPO 15 dakika ve RTO 4 saat yeterli mi?",
        "AI eğitimi için varsayılan anonim veri politikası ve istisna onay mercii kim olacak?",
        "Pilot bütçesi, üç yıllık TCO tavanı ve sağlayıcı çıkış bütçesi nedir?",
        "Teknik lider, veri koruma sorumlusu ve model yayın onay merciinin isimleri kimlerdir?",
    ):
        add_numbered(doc, text, num_decimal)

    doc.add_paragraph("13. Sonuç", style="Heading 1")
    add_body_paragraph(doc, "OY Dashboard için doğru yatırım, başlangıçta en büyük altyapıyı satın almak değil; veri büyüdükçe genişleyebilen ve her aşamada hukuki/teknik kanıt üreten bir işletim modeli kurmaktır. Bugünkü kamuya açık hizmet kapsamıyla Türk Telekom tek sağlayıcıda Türkiye içi veri ve GPU ihtiyacını en bütünlüklü karşılayan pilot adayıdır. Turkcell Bulut, yönetilen platform zenginliği ve geniş veri merkezi ağı nedeniyle kararın güçlü ikinci ayağıdır; NGN ise operatör bağımsız alternatif olarak rekabetçi teklif turunda tutulmalıdır.")
    add_body_paragraph(doc, "Önerinin başarı şartı sağlayıcı isminden çok sözleşme ve uygulama disiplinidir. Veri lokasyonu, destek erişimi, anahtar yönetimi, geri yükleme, AI veri seti onayı ve çıkış planı ölçülebilir hâle getirilmeden gerçek sağlık verisi canlı ortama alınmamalıdır.")
    add_callout(doc, "İlk yönetim kararı", "Üç sağlayıcılı RFP'nin başlatılması ve Türk Telekom üzerinde sentetik veriyle 30 günlük teknik pilot için bütçe/onay verilmesi.", fill=PALE_GREEN, accent=GREEN)

    page_break(doc)

    doc.add_paragraph("Ek A - Temel Tanımlar", style="Heading 1")
    glossary = [
        ("Açık rıza", "Belirli bir konuya ilişkin, bilgilendirmeye dayanan ve özgür iradeyle açıklanan rıza."),
        ("Anonimleştirme", "Verinin başka verilerle eşleştirilse dahi hiçbir şekilde belirli veya belirlenebilir kişiyle ilişkilendirilemeyecek hâle getirilmesi."),
        ("Pseudonimleştirme", "Doğrudan tanımlayıcıların ayrı tutulup kayıtların kodla temsil edilmesi. Ek bilgiyle kişi bulunabildiği için veri hâlâ kişisel veridir."),
        ("Özel nitelikli kişisel veri", "Sağlık, biyometrik ve genetik veri dahil Kanunda daha yüksek koruma verilen veri kategorileri."),
        ("Veri sorumlusu", "Kişisel verinin işleme amaç ve vasıtalarını belirleyen gerçek veya tüzel kişi."),
        ("Veri işleyen", "Veri sorumlusunun verdiği yetkiye dayanarak onun adına kişisel veri işleyen taraf; örneğin bulut sağlayıcısı."),
        ("Veri yerleşimi", "Verinin saklandığı ve işlendiği fiziksel ülke/bölgeye ilişkin kural ve teknik kontrol."),
        ("Veri egemenliği", "Verinin bulunduğu ülkenin hukuku, idari erişim koşulları ve operasyonel kontrolünün birlikte değerlendirilmesi."),
        ("RPO", "Bir kesintide kabul edilebilir en yüksek veri kaybı aralığı. RPO 15 dakika, son 15 dakikanın kaybedilebileceği anlamına gelir."),
        ("RTO", "Bir kesintiden sonra hizmetin yeniden çalışması için hedeflenen azami süre."),
        ("Felaket kurtarma (DR)", "Birincil ortamın kaybında farklı fiziksel bölgede hizmeti ve veriyi geri getiren plan, kapasite ve tatbikat bütünü."),
        ("Immutable/WORM yedek", "Belirli süre boyunca değiştirilemeyen veya silinemeyen, fidye yazılımına dayanıklı yedek kopya."),
        ("KMS/HSM", "Şifreleme anahtarlarının yaşam döngüsünü yöneten servis / anahtarları fiziksel güvenlikli donanımda koruyan modül."),
        ("AI eğitim alanı", "Eğitim verisi ve GPU kaynaklarının üretimden ayrıldığı, ağ çıkışı ve erişimi kısıtlı güvenlik bölgesi."),
        ("Veri seti kartı", "Bir eğitim setinin amacı, kaynağı, kapsamı, sınırlılığı, hukuki dayanağı, sürümü ve kalite/risk kayıtlarını açıklayan belge."),
        ("Model kartı", "Model amacı, eğitim yaklaşımı, performansı, riskleri, kullanım sınırları ve onay durumunu açıklayan kayıt."),
    ]
    add_table(doc, ["Terim", "Tanım"], glossary, [2350, 7010], header_fill=LIGHT_BLUE, font_size=8.8)

    doc.add_paragraph("Ek B - Kaynaklar", style="Heading 1")
    sources = [
        ("K1", "Kişisel Verileri Koruma Kurumu, Özel Nitelikli Kişisel Verilerin İşlenmesine İlişkin Rehber.", "https://www.kvkk.gov.tr/Icerik/8183/Ozel-Nitelikli-Kisisel-Verilerin-Islenmesine-Iliskin-Rehber"),
        ("K2", "KVKK Kurulunun 31.01.2018 tarihli ve 2018/10 sayılı özel nitelikli veri güvenliği kararı.", "https://www.kvkk.gov.tr/Icerik/4110/2018-10"),
        ("K3", "Kişisel Verileri Koruma Kurumu, Yapay Zekâ Alanında Kişisel Verilerin Korunmasına Dair Tavsiyeler.", "https://kvkk.gov.tr/SharedFolderServer/CMSFiles/25a1162f-0e61-4a43-98d0-3e7d057ac31a.pdf"),
        ("K4", "Kişisel Verileri Koruma Kurumu, Kişisel veri ihlali bildirimi ve 72 saat uygulaması.", "https://www.kvkk.gov.tr/Icerik/8595/kamuoyu-duyurusu"),
        ("K5", "Kişisel Verileri Koruma Kurumu, Kişisel Verilerin Yurt Dışına Aktarılması Rehberi.", "https://www.kvkk.gov.tr/Icerik/8143/Kisisel-Verilerin-Yurt-Disina-Aktarilmasi-Rehberi"),
        ("K6", "Kişisel Verileri Koruma Kurumu, Veri Sorumlusu ve Veri İşleyen rehberi.", "https://www.kvkk.gov.tr/Icerik/4195/Veri-Sorumlusu-ve-Veri-Isleyen"),
        ("K7", "Kişisel Verileri Koruma Kurumu, 04.09.2025 tarihli VERBİS istisna kararının uygulama esasları.", "https://www.kvkk.gov.tr/Icerik/8577/kisisel-verileri-koruma-kurulunun-04-09-2025-tarihli-ve-2025-1572-sayili-kararinin-uygulama-esaslarina-iliskin-kamuoyu-duyurusu"),
        ("K8", "Türk Telekom, Türkiye Genelinde Güvenli Veri Merkezleri ve Altyapı Hizmetleri.", "https://kurumsal.turktelekom.com.tr/bilisim-teknolojileri/veri-merkezi-ve-bulut/veri-merkezi-hizmetleri/veri-merkezlerimiz"),
        ("K9", "Türk Telekom, GPU'lu Sunucu Hizmetleri.", "https://kurumsal.turktelekom.com.tr/bilisim-teknolojileri/veri-merkezi-ve-bulut/sanallastirma-cozumleri/gpulu-sunucu-hizmetleri"),
        ("K10", "Turkcell, Veri Merkezi Çözümleri.", "https://www.turkcell.com.tr/kurumsal/dijital-is-servisleri/veri-merkezi-data-center"),
        ("K11", "Turkcell, Bulut Servisleri.", "https://www.turkcell.com.tr/kurumsal/dijital-is-servisleri/bulut-depolama"),
        ("K12", "NGN, Star of Bosphorus Veri Merkezi ve NGN Cloud duyuruları.", "https://www.ngn.com.tr/tr/star-of-bosphorus-veri-merkezi/"),
        ("K13", "Amazon Web Services, İstanbul Local Zone genel kullanıma açılış duyurusu, 20 Mayıs 2026.", "https://aws.amazon.com/about-aws/whats-new/2026/05/aws-local-zones-istanbul-turkiye/"),
        ("K14", "Amazon Web Services, EBS Local Snapshots ve veri yerleşimi açıklaması.", "https://docs.aws.amazon.com/pdfs/ebs/latest/userguide/ebs-ug.pdf"),
        ("K15", "Google Cloud, Türkiye'de yeni Cloud Region planı, 20 Kasım 2025.", "https://cloud.google.com/blog/products/infrastructure/new-google-cloud-region-coming-to-turkiye"),
        ("K16", "Turkcell, 2025 Form 20-F; Google Cloud Türkiye bölgesi öngörülen takvimi.", "https://s.turkcell.com.tr/SiteAssets/Hakkimizda/yatirimci-iliskileri/documents/pdf/20F2025.pdf"),
        ("K17", "Kişisel Verileri Koruma Kurumu, Kişisel Verilerin Silinmesi, Yok Edilmesi veya Anonim Hale Getirilmesi Rehberi.", "https://kvkk.gov.tr/SharedFolderServer/CMSFiles/bc1cb353-ef85-4e58-bb99-3bba31258508.pdf"),
    ]
    for key, label, url in sources:
        p = doc.add_paragraph(style="Source Text")
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        p.paragraph_format.keep_together = True
        r1 = p.add_run(f"[{key}] ")
        set_run_font(r1, size=9, color=NAVY, bold=True)
        r2 = p.add_run(label + " ")
        set_run_font(r2, size=9, color=INK)
        add_hyperlink(p, "Kaynağa erişim", url)

    p = doc.add_paragraph(style="Source Text")
    p.paragraph_format.space_before = Pt(10)
    r = p.add_run("Araştırma ve kaynak erişim tarihi: 28 Ağustos 2026. Sağlayıcı hizmet kapsamı, fiyatı ve sözleşme koşulları değişebileceğinden satın alma öncesinde güncel teklif ve kanıt istenmelidir.")
    set_run_font(r, size=9, color=MUTED, italic=True)

    # Avoid orphan headings and set table rows to repeat where appropriate.
    for paragraph in doc.paragraphs:
        if paragraph.style and paragraph.style.name.startswith("Heading"):
            paragraph.paragraph_format.keep_with_next = True
            paragraph.paragraph_format.keep_together = True

    doc.save(OUT_FILE)
    return OUT_FILE


if __name__ == "__main__":
    path = build_document()
    print(path)
